#!/usr/bin/env python3
"""Contract workflow for the maritime-service skill."""

from __future__ import annotations

import argparse
import json
import os
import re
import shutil
import sys
from copy import deepcopy
from dataclasses import dataclass
from datetime import date, datetime
from decimal import Decimal, InvalidOperation
from pathlib import Path
from typing import Any

from docx import Document
from openpyxl import Workbook, load_workbook


SCRIPT_DIR = Path(__file__).resolve().parent
SKILL_ROOT = SCRIPT_DIR.parent
ASSETS_DIR = SKILL_ROOT / "assets" / "contract_templates"
EXAMPLES_DIR = SKILL_ROOT / "examples" / "contract_requests"
EXAMPLE_WORKBOOK_DIR = SKILL_ROOT / "examples" / "workbooks"
DATA_ROOT = Path(os.environ.get("AMS_DATA_ROOT", str(SKILL_ROOT / "output"))).resolve()
OUTPUT_DIR = DATA_ROOT / "contracts"
REPORT_DIR = DATA_ROOT / "reports"
REGISTRY_PATH = SCRIPT_DIR / "contract_template_registry.json"

MONTHS_EN = {
    1: "JAN",
    2: "FEB",
    3: "MAR",
    4: "APR",
    5: "MAY",
    6: "JUN",
    7: "JUL",
    8: "AUG",
    9: "SEP",
    10: "OCT",
    11: "NOV",
    12: "DEC",
}

FIELD_ROWS = [
    ("template_key", "模板类型", "domestic_forwarder", "四选一：domestic_forwarder / domestic_logistics / overseas_astar / overseas_rta"),
    ("output_name", "输出文件名（可不填）", "mu-chuang-362-demo", "最终 .docx 文件名；不填则按合同号生成"),
    ("contract_no", "合同编号（可不填）", "MU CHANG 362 -202602-K--RONGTUA", "留空则按模板规则自动生成"),
    ("contract_date", "合同日期", "2026-02-04", "格式必须是 YYYY-MM-DD"),
    ("vessel_name", "船名", "MU CHUANG 362", "只填船名主体，不用自己加 MV"),
    ("vessel_particulars", "船舶规范（多行可用 | 分隔）", "BLT/OCT/2010 , CCS CLASS, CHINA FLAG, SDBC, DWT 35091MT ON ABT 10.27 SUMMER DRAFT|LOA179.88M, BEAM 28.8M, DEPTH 14.6M, GRT/NRT22382/11772|GRAIN/BALE CAPA:NO.1 6715.9 NO.2 9931.6,NO.3 9944.1 NO.4 9943.8 NO.5 8883.0|5H/5H, CRANE 4 X 30MT, ALL ABOUT AND WOG", "境内模板可填写多行；没有就留空"),
    ("cargo_summary_en", "货描英文", "ABT 18500MT SAINLESS STEEL COILS AND EMM IN TON BAG , 5% MOLCO", "建议直接填最终想展示的英文货描"),
    ("cargo_summary_cn", "货描中文", "约18500吨不锈钢卷和吨袋电解锰，±5%由租家决定", "建议直接填最终想展示的中文货描"),
    ("load_port_en", "装港英文完整句", "1SBP BAHODOPI, SULAWESI PROVINCE, INDONESIA   CHARTERERS' BERTH", "建议直接填完整合同句子"),
    ("load_port_cn", "装港中文完整句", "印尼苏拉威西BAHODOPI港一个安全泊位", "建议直接填完整合同句子"),
    ("discharge_port_en", "卸港英文完整句", "1SBP TANGJUNG LANGSAT + 1 SBP KELANG(WEST PORT), MALAYSIA   OWNERS' BERTH.", "建议直接填完整合同句子"),
    ("discharge_port_cn", "卸港中文完整句", "马来西亚巴鲁钢的一个安全泊位，和巴生西港的一个安全泊位", "建议直接填完整合同句子"),
    ("laycan_start", "LAYCAN 开始日期", "2026-02-02", "格式必须是 YYYY-MM-DD"),
    ("laycan_end", "LAYCAN 结束日期", "2026-02-05", "格式必须是 YYYY-MM-DD"),
    ("laycan_en", "LAYCAN 英文（可不填）", "2/FEB/2026-5/FEB/2026", "留空则自动生成"),
    ("laycan_cn", "LAYCAN 中文（可不填）", "2026年2月2日- 2026年2月5日受载", "留空则自动生成"),
    ("demurrage_rate_usd", "滞期费美元/天", "6500", "只填数字"),
    ("loading_rate_en", "装卸率英文", "LDG 4 WWD SHINC, DISCHG CQD. ONCE DEMURRAGE, ALWAYS ON DEMURRAGE", "建议直接填最终句子"),
    ("loading_rate_cn", "装卸率中文", "装港为4个晴天工作日，卸港当地港口通常的装卸速度。一旦滞期，永远滞期", "建议直接填最终句子"),
    ("freight_rate_usd_per_mt", "运费美元/吨", "17.65", "只填数字"),
    ("freight_terms_code", "运费条款代码", "FILO", "常见为 FILO / FIO"),
    ("freight_port_basis", "运费港口基数", "1/2", "例如 1/2、1/1"),
    ("freight_clause_note_en", "运费条款英文补充（可不填）", "", "例如 (AL-INGOT FIO)"),
    ("freight_clause_note_cn", "运费条款中文补充（可不填）", "", "例如 (铝锭FIO)"),
    ("bank_info_lines", "船东银行信息（可不填，多行用 | 分隔）", "", "留空则使用模板原文"),
    ("signature_charterer", "签字栏左侧公司名（可不填）", "", "留空则使用模板原文"),
    ("signature_owner", "签字栏右侧公司名（可不填）", "", "留空则使用模板原文"),
    ("signature_charterer_detail", "签字栏左侧联系人（可不填）", "", "留空则使用模板原文"),
    ("signature_owner_detail", "签字栏右侧联系人（可不填）", "", "留空则使用模板原文")
]

CARGO_HEADERS = [
    "loading_port",
    "discharge_port",
    "quantity",
    "unit",
    "cargo_name_en",
    "cargo_name_cn",
    "terms"
]


class ContractError(Exception):
    """Base contract workflow error."""


class ContractValidationError(ContractError):
    """Validation error with multiple user-facing messages."""

    def __init__(self, errors: list[str]) -> None:
        super().__init__("\n".join(errors))
        self.errors = errors


@dataclass
class RenderResult:
    document_path: Path
    summary_path: Path
    summary: dict[str, Any]


def load_registry() -> dict[str, dict[str, Any]]:
    registry = json.loads(REGISTRY_PATH.read_text(encoding="utf-8"))
    for item in registry.values():
        item["template_path"] = (ASSETS_DIR / item["template_file"]).resolve()
    return registry


def read_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def write_json(path: Path, payload: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def parse_date(value: Any, field_name: str) -> date:
    if isinstance(value, datetime):
        return value.date()
    if isinstance(value, date):
        return value
    text = str(value or "").strip()
    if not text:
        raise ContractValidationError([f"字段 `{field_name}` 不能为空，且必须是 YYYY-MM-DD 格式。"])
    try:
        return date.fromisoformat(text)
    except ValueError as exc:
        raise ContractValidationError([f"字段 `{field_name}` 不是合法日期：{text!r}。请使用 YYYY-MM-DD。"]) from exc


def parse_decimal(value: Any, field_name: str) -> Decimal:
    text = str(value or "").strip()
    if not text:
        raise ContractValidationError([f"字段 `{field_name}` 不能为空。"])
    try:
        return Decimal(text)
    except InvalidOperation as exc:
        raise ContractValidationError([f"字段 `{field_name}` 不是合法数字：{text!r}。"]) from exc


def compact_spaces(text: str) -> str:
    return re.sub(r"\s+", " ", text.strip())


def vessel_name_upper(value: str) -> str:
    text = compact_spaces(value)
    text = re.sub(r"^MV\s+", "", text, flags=re.IGNORECASE)
    return text.upper()


def display_vessel_name(value: str) -> str:
    text = compact_spaces(value)
    if text.upper().startswith("MV "):
        return text
    return f"MV {text}"


def ordinal_suffix(day: int) -> str:
    if 10 <= day % 100 <= 20:
        return "TH"
    return {1: "ST", 2: "ND", 3: "RD"}.get(day % 10, "TH")


def format_contract_date_en(value: date) -> str:
    return f"{value.day}{ordinal_suffix(value.day)} {MONTHS_EN[value.month]} {value.year}"


def format_laycan_en(start: date, end: date) -> str:
    return f"{start.day}/{MONTHS_EN[start.month]}/{start.year}-{end.day}/{MONTHS_EN[end.month]}/{end.year}"


def format_laycan_cn(start: date, end: date) -> str:
    return f"{start.year}年{start.month}月{start.day}日- {end.year}年{end.month}月{end.day}日受载"


def decimal_to_string(value: Decimal) -> str:
    normalized = value.normalize()
    text = format(normalized, "f")
    return text.rstrip("0").rstrip(".") if "." in text else text


def maybe_split_lines(value: Any) -> list[str]:
    if isinstance(value, list):
        return [str(item).strip() for item in value if str(item).strip()]
    text = str(value or "").strip()
    if not text:
        return []
    return [item.strip() for item in text.split("|") if item.strip()]


def unique_join(items: list[str], separator: str) -> str:
    seen: list[str] = []
    for item in items:
        item = compact_spaces(item)
        if item and item not in seen:
            seen.append(item)
    return separator.join(seen)


def derive_cargo_summary(request: dict[str, Any], rows: list[dict[str, Any]]) -> tuple[str, str]:
    cargo_summary_en = compact_spaces(str(request.get("cargo_summary_en", "")))
    cargo_summary_cn = compact_spaces(str(request.get("cargo_summary_cn", "")))
    if cargo_summary_en and cargo_summary_cn:
        return cargo_summary_en, cargo_summary_cn

    quantities = [parse_decimal(row["quantity"], "cargo_rows.quantity") for row in rows if str(row.get("quantity", "")).strip()]
    total_quantity = sum(quantities, start=Decimal("0"))
    unit = compact_spaces(str(rows[0].get("unit", "MT"))) if rows else "MT"
    cargo_names_en = unique_join([str(row.get("cargo_name_en", "")) for row in rows], ", ")
    cargo_names_cn = unique_join([str(row.get("cargo_name_cn", "")) for row in rows], "、")
    total_text = decimal_to_string(total_quantity)

    if not cargo_summary_en:
        cargo_summary_en = f"ABT {total_text}{unit} {cargo_names_en} , 5% MOLCO".strip()
    if not cargo_summary_cn:
        cargo_summary_cn = f"约{total_text}吨{cargo_names_cn}，±5%由租家决定"
    return cargo_summary_en, cargo_summary_cn


def derive_contract_no(request: dict[str, Any], template: dict[str, Any], contract_date: date) -> str:
    manual_value = compact_spaces(str(request.get("contract_no", "")))
    if manual_value:
        return manual_value

    pattern = template["contract_no_pattern"]
    vessel_key = vessel_name_upper(str(request.get("vessel_name", "")))
    return pattern.format(vessel_name_upper=vessel_key, yyyymm=f"{contract_date.year}{contract_date.month:02d}")


def derive_demurrage_texts(amount: Decimal) -> dict[str, str]:
    amount_text = decimal_to_string(amount)
    return {
        "out_of_laycan_clause_2_en": "(2) OUT OF LAYCAN OVER 2 DAYS BUT LESS THAN(INCLUDING)3 DAYS, BESIDES CLASUE (1), OWNER SHOULD PAY USD " + amount_text + "/DAY PDPR AS COMPENSATION FOR OVER 2DAYS PARTS;",
        "out_of_laycan_clause_3_en": "(3) OUT OF LAYCAN OVER 3 DAYS, BESIDES CLAUSE (1), OWNER SHOULD PAY USD " + amount_text + "/DAY PDPR AS COMPENSATION FM THE CANCELING DAY TILL NOR TENDERED;",
        "out_of_laycan_clause_2_cn": "(2)迟到2-3天（包括3天），除第(1)条外，超过2天的部分按比例收反向滞期费；",
        "out_of_laycan_clause_3_cn": "(3)迟到大于3天,除第(1)条外，所有的迟到时间都收反向滞期费；",
        "demurrage_clause_en": "DETENTION/DEMURRAGE/ IN AMOUNT USD " + amount_text + " PDPR/DHD.DETENTION ONLY APPLICABLE IN CASE OF VESSELS DELAY DUE TO THE FACT THAT CARGO AND/OR CARGO DOCUMENTS ARE NOT READY ON VESSEL’S ARRIVAL TO LOADING/DISCHARGE PORT.",
        "demurrage_clause_cn": "滞留费/滞期费/USD" + amount_text + "/天，不足一天按比例算，速遣为滞期的一半，滞留费适用于船舶抵港及递交NOR后因货物和/或货物单证未备妥而导致船舶延误。"
    }


def derive_freight_texts(rate: Decimal, terms_code: str, port_basis: str, note_en: str, note_cn: str) -> dict[str, str]:
    rate_text = decimal_to_string(rate)
    note_en_part = f" {note_en.strip()}" if note_en.strip() else ""
    note_cn_part = f"{note_cn.strip()}" if note_cn.strip() else ""
    discharge_count = port_basis.split("/")[-1] if "/" in port_basis else port_basis
    return {
        "freight_clause_en": f"USD{rate_text}/MT, {terms_code.strip()} TERMS BSS {port_basis.strip()}{note_en_part}, AND FRT SHALL BE CALCULATED AS PER B/L’S CARGO QTY.",
        "freight_clause_cn": f"海运费{rate_text}美元每吨， {terms_code.strip()}条款{note_cn_part}，一个装港{discharge_count}个卸港，以提单数量为结算依据。"
    }


def sanitize_filename(value: str) -> str:
    value = compact_spaces(value)
    value = re.sub(r"[\\/:*?\"<>|]+", "_", value)
    return value or "generated-contract"


def set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = text


def set_paragraph_multiline(paragraph, lines: list[str]) -> None:
    if not lines:
        set_paragraph_text(paragraph, "")
        return
    if not paragraph.runs:
        paragraph.text = "\n".join(lines)
        return
    paragraph.runs[0].text = "\n".join(lines)
    for run in paragraph.runs[1:]:
        run.text = ""


def set_cell_text(cell, text: str) -> None:
    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]._element
        p.getparent().remove(p)
    if cell.paragraphs:
        set_paragraph_text(cell.paragraphs[0], text)
    else:
        cell.text = text


def rewrite_table_rows(table, rows: list[dict[str, Any]], columns: list[str]) -> None:
    sample_row = deepcopy(table.rows[0]._tr)
    table_element = table._tbl
    for row in list(table.rows):
        table_element.remove(row._tr)
    for payload in rows:
        new_row_tr = deepcopy(sample_row)
        table_element.append(new_row_tr)
        current_row = table.rows[-1]
        for column_index, key in enumerate(columns):
            value = compact_spaces(str(payload.get(key, "")))
            set_cell_text(current_row.cells[column_index], value)


def load_json_request(path: Path) -> dict[str, Any]:
    return read_json(path)


def load_workbook_request(path: Path) -> dict[str, Any]:
    workbook = load_workbook(path, data_only=True)
    if "头信息" not in workbook.sheetnames or "货物明细" not in workbook.sheetnames:
        raise ContractValidationError(["工作簿必须包含 `头信息` 和 `货物明细` 两个工作表。"])

    header_sheet = workbook["头信息"]
    cargo_sheet = workbook["货物明细"]
    request: dict[str, Any] = {}

    for row in header_sheet.iter_rows(min_row=2, values_only=True):
        key = str(row[0] or "").strip()
        value = row[3] if len(row) > 3 else None
        if key and value not in (None, ""):
            request[key] = value

    cargo_rows: list[dict[str, Any]] = []
    for row in cargo_sheet.iter_rows(min_row=2, values_only=True):
        if not any(item not in (None, "") for item in row):
            continue
        payload = {header: row[index] if index < len(row) else "" for index, header in enumerate(CARGO_HEADERS)}
        if not str(payload.get("loading_port", "")).strip():
            continue
        cargo_rows.append(payload)
    request["cargo_rows"] = cargo_rows
    return request


def validate_request(request: dict[str, Any], registry: dict[str, dict[str, Any]]) -> dict[str, Any]:
    errors: list[str] = []
    template_key = compact_spaces(str(request.get("template_key", "")))
    if not template_key:
        errors.append("必须填写 `template_key`。")
    elif template_key not in registry:
        errors.append(f"`template_key` 不合法：{template_key!r}。")

    required_text_fields = [
        "contract_date",
        "vessel_name",
        "load_port_en",
        "load_port_cn",
        "discharge_port_en",
        "discharge_port_cn",
        "laycan_start",
        "laycan_end",
        "demurrage_rate_usd",
        "loading_rate_en",
        "loading_rate_cn",
        "freight_rate_usd_per_mt",
        "freight_terms_code",
        "freight_port_basis"
    ]
    for field_name in required_text_fields:
        if not str(request.get(field_name, "")).strip():
            errors.append(f"字段 `{field_name}` 不能为空。")

    cargo_rows = request.get("cargo_rows", [])
    if not isinstance(cargo_rows, list) or not cargo_rows:
        errors.append("至少需要 1 行 `cargo_rows` 货物明细。")
    else:
        for index, row in enumerate(cargo_rows, start=1):
            for key in ["loading_port", "discharge_port", "quantity", "unit", "cargo_name_en"]:
                if not str(row.get(key, "")).strip():
                    errors.append(f"`cargo_rows` 第 {index} 行缺少 `{key}`。")

    if errors:
        raise ContractValidationError(errors)

    template = registry[template_key]
    contract_date = parse_date(request["contract_date"], "contract_date")
    laycan_start = parse_date(request["laycan_start"], "laycan_start")
    laycan_end = parse_date(request["laycan_end"], "laycan_end")
    demurrage_rate = parse_decimal(request["demurrage_rate_usd"], "demurrage_rate_usd")
    freight_rate = parse_decimal(request["freight_rate_usd_per_mt"], "freight_rate_usd_per_mt")

    cargo_summary_en, cargo_summary_cn = derive_cargo_summary(request, cargo_rows)
    normalized = dict(request)
    normalized["template_key"] = template_key
    normalized["template"] = template
    normalized["contract_date_obj"] = contract_date
    normalized["laycan_start_obj"] = laycan_start
    normalized["laycan_end_obj"] = laycan_end
    normalized["contract_no"] = derive_contract_no(request, template, contract_date)
    normalized["contract_date_en"] = compact_spaces(str(request.get("contract_date_en", ""))) or format_contract_date_en(contract_date)
    normalized["laycan_en"] = compact_spaces(str(request.get("laycan_en", ""))) or format_laycan_en(laycan_start, laycan_end)
    normalized["laycan_cn"] = compact_spaces(str(request.get("laycan_cn", ""))) or format_laycan_cn(laycan_start, laycan_end)
    normalized["vessel_name_display"] = display_vessel_name(str(request["vessel_name"]))
    normalized["vessel_particular_lines"] = maybe_split_lines(request.get("vessel_particulars", request.get("vessel_particular_lines", [])))
    normalized["cargo_summary_en"] = cargo_summary_en
    normalized["cargo_summary_cn"] = cargo_summary_cn
    normalized["demurrage_texts"] = derive_demurrage_texts(demurrage_rate)
    normalized["freight_texts"] = derive_freight_texts(
        freight_rate,
        str(request["freight_terms_code"]),
        str(request["freight_port_basis"]),
        str(request.get("freight_clause_note_en", "")),
        str(request.get("freight_clause_note_cn", ""))
    )
    normalized["bank_info_lines"] = maybe_split_lines(request.get("bank_info_lines", []))
    normalized["signature_charterer"] = compact_spaces(str(request.get("signature_charterer", "")))
    normalized["signature_owner"] = compact_spaces(str(request.get("signature_owner", "")))
    normalized["signature_charterer_detail"] = compact_spaces(str(request.get("signature_charterer_detail", "")))
    normalized["signature_owner_detail"] = compact_spaces(str(request.get("signature_owner_detail", "")))
    normalized["output_name"] = compact_spaces(str(request.get("output_name", "")))
    return normalized


def apply_optional_bank_info(doc: Document, template: dict[str, Any], request: dict[str, Any]) -> None:
    bank_info_lines = request.get("bank_info_lines", [])
    if not bank_info_lines:
        return
    indices = template["paragraphs"].get("bank_info", [])
    for idx, line in zip(indices, bank_info_lines):
        set_paragraph_text(doc.paragraphs[idx], line)
    for idx in indices[len(bank_info_lines):]:
        set_paragraph_text(doc.paragraphs[idx], "")


def apply_optional_signature_overrides(doc: Document, template: dict[str, Any], request: dict[str, Any]) -> None:
    table_index = template.get("signature_table_index")
    if table_index is None:
        return
    table = doc.tables[table_index]
    if request.get("signature_charterer"):
        set_cell_text(table.rows[1].cells[0], request["signature_charterer"])
    if request.get("signature_owner"):
        set_cell_text(table.rows[1].cells[1], request["signature_owner"])
    if request.get("signature_charterer_detail"):
        set_cell_text(table.rows[2].cells[0], request["signature_charterer_detail"])
    if request.get("signature_owner_detail"):
        set_cell_text(table.rows[2].cells[1], request["signature_owner_detail"])


def render_contract(request: dict[str, Any], output_dir: Path, summary_dir: Path | None = None) -> RenderResult:
    template = request["template"]
    doc = Document(template["template_path"])
    paragraphs = template["paragraphs"]
    output_dir = output_dir.resolve()
    summary_dir = (summary_dir or REPORT_DIR).resolve()

    set_paragraph_text(doc.paragraphs[paragraphs["contract_no"]], f"CONTRACT NO.: {request['contract_no']}")
    set_paragraph_text(doc.paragraphs[paragraphs["date"]], f"DATED: {request['contract_date_en']}")
    set_paragraph_text(doc.paragraphs[paragraphs["vessel_name"]], request["vessel_name_display"])

    if "vessel_particulars" in paragraphs:
        set_paragraph_multiline(doc.paragraphs[paragraphs["vessel_particulars"]], request["vessel_particular_lines"])

    set_paragraph_text(doc.paragraphs[paragraphs["cargo_summary_en"]], request["cargo_summary_en"])
    set_paragraph_text(doc.paragraphs[paragraphs["cargo_summary_cn"]], request["cargo_summary_cn"])
    set_paragraph_text(doc.paragraphs[paragraphs["load_port_en"]], compact_spaces(str(request["load_port_en"])))
    set_paragraph_text(doc.paragraphs[paragraphs["load_port_cn"]], compact_spaces(str(request["load_port_cn"])))
    set_paragraph_text(doc.paragraphs[paragraphs["discharge_port_en"]], compact_spaces(str(request["discharge_port_en"])))
    set_paragraph_text(doc.paragraphs[paragraphs["discharge_port_cn"]], compact_spaces(str(request["discharge_port_cn"])))
    set_paragraph_text(doc.paragraphs[paragraphs["laycan_en"]], request["laycan_en"])
    set_paragraph_text(doc.paragraphs[paragraphs["laycan_cn"]], request["laycan_cn"])

    for key in ["out_of_laycan_clause_2_en", "out_of_laycan_clause_3_en", "out_of_laycan_clause_2_cn", "out_of_laycan_clause_3_cn", "demurrage_clause_en", "demurrage_clause_cn"]:
        set_paragraph_text(doc.paragraphs[paragraphs[key]], request["demurrage_texts"][key])
    for key in ["freight_clause_en", "freight_clause_cn"]:
        set_paragraph_text(doc.paragraphs[paragraphs[key]], request["freight_texts"][key])

    set_paragraph_text(doc.paragraphs[paragraphs["loading_rate_en"]], compact_spaces(str(request["loading_rate_en"])))
    set_paragraph_text(doc.paragraphs[paragraphs["loading_rate_cn"]], compact_spaces(str(request["loading_rate_cn"])))

    if template["layout"] == "domestic_table":
        table = doc.tables[template["cargo_table_index"]]
        rewrite_table_rows(table, request["cargo_rows"], template["cargo_table_columns"])

    apply_optional_bank_info(doc, template, request)
    apply_optional_signature_overrides(doc, template, request)

    output_dir.mkdir(parents=True, exist_ok=True)
    summary_dir.mkdir(parents=True, exist_ok=True)
    base_name = sanitize_filename(request["output_name"] or request["contract_no"])
    document_path = output_dir / f"{base_name}.docx"
    summary_path = summary_dir / f"{base_name}.json"
    doc.save(document_path)

    summary = {
        "request_id": request.get("request_id", ""),
        "template_key": request["template_key"],
        "document_path": str(document_path),
        "contract_no": request["contract_no"],
        "contract_date": request["contract_date_obj"].isoformat(),
        "vessel_name": request["vessel_name_display"],
        "cargo_rows": request["cargo_rows"]
    }
    write_json(summary_path, summary)
    return RenderResult(document_path=document_path, summary_path=summary_path, summary=summary)


def extract_doc_text(path: Path) -> str:
    doc = Document(path)
    lines = [compact_spaces(paragraph.text) for paragraph in doc.paragraphs if compact_spaces(paragraph.text)]
    for table in doc.tables:
        for row in table.rows:
            lines.append(" | ".join(compact_spaces(cell.text) for cell in row.cells))
    return "\n".join(lines)


def create_workbook_template(output_path: Path, example_request: dict[str, Any] | None = None) -> Path:
    workbook = Workbook()
    info_sheet = workbook.active
    info_sheet.title = "说明"
    info_sheet["A1"] = "这是合同填写模板。请先看 docs/01-小学生级别使用教程.md。"
    info_sheet["A2"] = "真正要填写的是“头信息”工作表的 D 列，以及“货物明细”工作表。"
    info_sheet["A3"] = "日期统一使用 YYYY-MM-DD。多行文本可用 | 分隔。"
    info_sheet["A4"] = "模板类型只能填：domestic_forwarder / domestic_logistics / overseas_astar / overseas_rta"

    head_sheet = workbook.create_sheet("头信息")
    head_sheet.append(["字段key", "中文说明", "示例值", "实际值"])
    example_request = example_request or {}
    for field_key, label, sample, hint in FIELD_ROWS:
        actual = ""
        if field_key in example_request:
            actual_value = example_request[field_key]
            actual = "|".join(str(item) for item in actual_value) if isinstance(actual_value, list) else str(actual_value)
        head_sheet.append([field_key, f"{label} / {hint}", sample, actual])

    cargo_sheet = workbook.create_sheet("货物明细")
    cargo_sheet.append(CARGO_HEADERS)
    for row in example_request.get("cargo_rows", []):
        cargo_sheet.append([row.get(header, "") for header in CARGO_HEADERS])

    output_path.parent.mkdir(parents=True, exist_ok=True)
    workbook.save(output_path)
    return output_path


def build_example_workbooks() -> list[Path]:
    built: list[Path] = []
    for example_path in sorted(EXAMPLES_DIR.glob("*.json")):
        request = load_json_request(example_path)
        output_path = EXAMPLE_WORKBOOK_DIR / f"{example_path.stem}.xlsx"
        built.append(create_workbook_template(output_path, request))
    blank_path = EXAMPLE_WORKBOOK_DIR / "blank-contract-template.xlsx"
    built.append(create_workbook_template(blank_path, None))
    return built


def build_examples(registry: dict[str, dict[str, Any]]) -> list[RenderResult]:
    results: list[RenderResult] = []
    EXAMPLE_WORKBOOK_DIR.mkdir(parents=True, exist_ok=True)
    for example_path in sorted(EXAMPLES_DIR.glob("*.json")):
        request = validate_request(load_json_request(example_path), registry)
        results.append(render_contract(request, OUTPUT_DIR))
    build_example_workbooks()
    return results


def verify_examples(registry: dict[str, dict[str, Any]]) -> list[str]:
    messages: list[str] = []
    verification_dir = OUTPUT_DIR / "_verification"
    if verification_dir.exists():
        shutil.rmtree(verification_dir)
    verification_dir.mkdir(parents=True, exist_ok=True)

    for example_path in sorted(EXAMPLES_DIR.glob("*.json")):
        request = validate_request(load_json_request(example_path), registry)
        result = render_contract(request, verification_dir)
        text = extract_doc_text(result.document_path)
        checks = [request["contract_no"], request["contract_date_en"], request["vessel_name_display"], request["cargo_summary_en"]]
        for row in request["cargo_rows"][:2]:
            checks.append(compact_spaces(str(row["loading_port"])))
            checks.append(compact_spaces(str(row["discharge_port"])))
        missing = [value for value in checks if value and value not in text]
        if missing:
            raise ContractError(f"{example_path.name} 校验失败，生成文件缺少这些内容：{missing}")
        messages.append(f"{example_path.name} -> 通过")
    return messages


def print_result(result: RenderResult) -> None:
    print(f"[OK] 合同已生成: {result.document_path}")
    print(f"[OK] 摘要已生成: {result.summary_path}")
    print(f"[INFO] 合同编号: {result.summary['contract_no']}")


def load_request_from_args(args: argparse.Namespace) -> dict[str, Any]:
    if args.command == "from-json":
        return load_json_request(Path(args.input))
    if args.command == "from-workbook":
        return load_workbook_request(Path(args.input))
    raise ContractError(f"不支持的命令：{args.command}")


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description=__doc__)
    subparsers = parser.add_subparsers(dest="command", required=True)

    parser_json = subparsers.add_parser("from-json", help="从 JSON 请求生成合同")
    parser_json.add_argument("--input", required=True, help="JSON 请求文件路径")
    parser_json.add_argument("--output-dir", default=str(OUTPUT_DIR), help="输出目录")
    parser_json.add_argument("--summary-dir", default=str(REPORT_DIR), help="JSON 摘要输出目录")

    parser_workbook = subparsers.add_parser("from-workbook", help="从 Excel 工作簿生成合同")
    parser_workbook.add_argument("--input", required=True, help="Excel 工作簿路径")
    parser_workbook.add_argument("--output-dir", default=str(OUTPUT_DIR), help="输出目录")
    parser_workbook.add_argument("--summary-dir", default=str(REPORT_DIR), help="JSON 摘要输出目录")

    parser_make_workbook = subparsers.add_parser("make-workbook-template", help="生成 Excel 填写模板")
    parser_make_workbook.add_argument("--output", required=True, help="输出的 xlsx 路径")
    parser_make_workbook.add_argument("--example-json", help="可选：使用某个 JSON 示例预填实际值")

    subparsers.add_parser("build-examples", help="生成所有示例合同和示例 Excel")
    subparsers.add_parser("verify-examples", help="生成并校验所有示例合同")
    return parser


def main(argv: list[str] | None = None) -> int:
    sys.stdout.reconfigure(encoding="utf-8")
    parser = build_parser()
    args = parser.parse_args(argv)
    registry = load_registry()

    try:
        if args.command == "make-workbook-template":
            example_request = load_json_request(Path(args.example_json)) if args.example_json else None
            output_path = create_workbook_template(Path(args.output), example_request)
            print(f"[OK] Excel 模板已生成: {output_path}")
            return 0

        if args.command == "build-examples":
            results = build_examples(registry)
            print("[OK] 示例合同已全部生成。")
            for result in results:
                print_result(result)
            return 0

        if args.command == "verify-examples":
            messages = verify_examples(registry)
            print("[OK] 示例校验通过：")
            for message in messages:
                print(f" - {message}")
            return 0

        request = load_request_from_args(args)
        normalized_request = validate_request(request, registry)
        result = render_contract(
            normalized_request,
            Path(args.output_dir),
            Path(args.summary_dir),
        )
        print_result(result)
        return 0
    except ContractValidationError as exc:
        print("[ERROR] 输入校验失败：")
        for item in exc.errors:
            print(f" - {item}")
        return 2
    except ContractError as exc:
        print(f"[ERROR] {exc}")
        return 3


if __name__ == "__main__":
    raise SystemExit(main())
